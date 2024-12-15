import re
from docx.oxml import OxmlElement


# Функция для замены шаблонов формул в тексте
def replace_formulas_with_math_objects(doc, text, formulas):
    # Регулярное выражение для поиска шаблонов формул
    pattern = re.compile(r'%<([^>]+)>%')
    last_end = 0

    for match in pattern.finditer(text):
        # Добавляем текст до формулы
        doc.add_paragraph(text[last_end:match.start()])

        formula_id = match.group(1)
        latex_formula = formulas.get(formula_id, None)

        if latex_formula:
            add_latex_as_math(doc, latex_formula)
        else:
            # Если формула не найдена, оставляем оригинальный текст
            doc.add_paragraph(match.group(0))

        last_end = match.end()

    # Добавляем остаток текста после последней формулы
    if last_end < len(text):
        doc.add_paragraph(text[last_end:])


# Функция для добавления формулы в OMML формате
def add_latex_as_math(doc, latex_formula):
    # Создаем пустой параграф
    p = doc.add_paragraph()

    # Добавляем OMML (Office Math Markup Language) элемент
    math_element = OxmlElement('m:oMathPara')
    oMath = OxmlElement('m:oMath')
    math_element.append(oMath)
    p._element.append(math_element)

    # Внутри OMML элемент заполняется содержимым
    math_run = OxmlElement('m:r')
    math_text = OxmlElement('m:t')
    math_text.text = latex_formula  # Добавляем LaTeX формулу как текст
    math_run.append(math_text)
    oMath.append(math_run)
